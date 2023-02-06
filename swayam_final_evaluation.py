from tkinter import*
from tkinter import filedialog
from tkinter import font 
import tkinter.messagebox as tmsg
from tkinter import colorchooser
from PIL import Image, ImageTk
import docx
from docx.shared import Pt
import PyPDF2
import os
import pdfplumber


root=Tk()
root.title('Untitled -Notepad')
root.geometry("1200x660")

#set variable for open file 
global file_global
file_global=False

global selected
selected=False

#create function for new file
def new_f():
    #delet previous text
    TextArea.delete("1.0",END)
    #update title and status bar
    root.title('Untitled -Notepad')
    status_bar.config(text="New File        ")

    #set variable for open file 
    global file_global
    file_global=False

#create function for open pdf
# def open_pdf():
#     TextArea.delete("1.0",END)
#     file=filedialog.askopenfilename(title="Select a PDF",filetypes=(("PDF File","*.pdf"),("All Files","*.*")))
#     if file:
#         global file_global
#         file_global=file 
#         #open the PDF file
#         pdf_file=PyPDF2.PdfFileReader(file)
#         #Select a page to read
#         page=pdf_file.getPage(0)
#         #get the content of the pages 
#         content=page.extractText()
#         #Add the content to TextArea
#         TextArea.insert(1.0,content)

#         filename=file
#         #update status bar 
#         status_bar.config(text=f"{filename}        ")
#         #update title
#         #filename=filename.replace("C:/Users/Owners/Documents/","")
#         root.title(f"{filename} -Notepad")

# #create function for docx file
# def open_docx():
#     TextArea.delete("1.0",END)
#     file=filedialog.openfileename(title="select a DOCX file",filetypes=(("Docx file","*.docx"),("All File","*.*")))

#     global file_global
#     file_global=file 
#     if file:
#         #create an instance of a word
#         doc=docx.Document()
#         #add a heading of level 0
#         doc.add_heading(file,0)

#         #add a pargraph
#         doc_para= doc.add_paragraph().add_run(END)

#         filename=file
#         #update status bar 
#         status_bar.config(text=f"{filename}        ")
#         #update title
#         #filename=filename.replace("C:/Users/Owners/Documents/","")
#         root.title(f"{filename} -Notepad")
#         #open the docx file

#creat function for Open file
def open_f():
    #delet previous text
    #TextArea.delete("1.0",END)
    file=filedialog.askopenfilename(initialdir="C:/Users/Owners/Documents/ ",filetypes=[("Text Document",".txt"),("Docx File",".docx"),("PDF File",".pdf"),("ALL files",".")])
    #
    #check to see there is a file name 
    if file:
        #make file name globale so we can exsses this other
        global file_global
        file_global=file 

    filename=file
    #show file adress in status bar (update statusbar)
    status_bar.config(text=f"{filename}        ")
    #update title with newfile name 
    filename=filename.replace("C:/Users/Owners/Documents/"," ")
    root.title(f"{filename} -Notepad")

    #open file 
    split_tup=os.path.splitext(file)
    file_extension=split_tup[1]

    textToInsert=""
    #for text file
    if (file_extension==".txt"):
        TextArea.delete("1.0",END)
        file=open(file,'r')
        TextArea.insert(END,file.read())
        #file.close()

    elif (file_extension==".pdf"):
        pdfFile=open(file,'rb')
        readpdf=PyPDF2.PdfReader(pdfFile)
        totalpages=len(readpdf.pages)
        #global textToInsert
        with pdfplumber.open(file) as f:
            for j in range(totalpages):
                data=f.pages[j]
                textToInsert += data.extract_text()
                TextArea.delete("1.0",END)
                TextArea.insert(END,textToInsert)
    
    elif (file_extension==".docx"):
        doc = docx.Document(file)
        content = doc.paragraphs
        TextArea.delete("1.0",END)
        for para in content:
            TextArea.insert(END,para.text)
            TextArea.insert(END,"\n")

#create function for save as file
def save_as_f():
    file=filedialog.asksaveasfilename(defaultextension=".",initialdir="C:/Users/Owners/Documents/",title="Save File",filetypes=[("ALL files","*.*"),("Text Document","*.txt"),("Python File","*.py"),("C File","*.c")])
    if file:
        filename=file
        #update status bar 
        status_bar.config(text=f"Saved: {filename}        ")
        #update title
        filename=filename.replace("C:/Users/Owners/Documents/","")
        root.title(f"{filename} -Notepad")

        #save file
        file=open(file,'w')
        file.write(TextArea.get(1.0,END))
        file.close()

#create function for save docx
def save_as_docx():
    file=filedialog.asksaveasfilename(defaultextension=".",initialdir="C:/Users/Owners/Documents/",title="Save File",filetypes=[("ALL files","."),("Docx File", ".docx")])

    st = TextArea.get(1.0, END)
    # print(type(str2))
    doc = docx.Document()
    # doc.add_heading('Rushi', 0)
    para = doc.add_paragraph().add_run(st)
    para.font.size = Pt(16)
    doc.save(file)

#create function for save
def save_f():
    global file_global
    
    if file_global:
        #save file
        file=open(file_global,'w')
        file.write(TextArea.get(1.0,END))
        file.close()
        #update status
        status_bar.config(text=f"Saved: {file_global}        ")
    else:
        save_as_f()

#create function for cut  text 
def cut_f():
    TextArea.event_generate(("<<Cut>>"))

#create function for copy  text 
def copy_f():
    TextArea.event_generate(("<<Copy>>"))

#create function for paste text 
def paste_f():
    TextArea.event_generate(("<<Paste>>"))

#create function for select all
def select_all_f(e):
    #add sel tag to selct all select
    TextArea.tag_add('sel','1.0',END)

#create function for clear all
def clear_f():
    TextArea.delete(1.0,END)

#create function for help 
def about_f():
    ab=tmsg.showinfo("help","this a help box for Notepad")

#crete a function for bold text 
def bold_it():
    #creat our font 
    bold_font= font.Font(TextArea,TextArea.cget("font"))
    bold_font.config(weight="bold")

    #configur a tag 
    TextArea.tag_configure("bold",font=bold_font)

    #define current tag
    current_tag = TextArea.tag_names("sel.first")
    #if stetment to see if tag has been bold 
    if "bold" in current_tag:
        TextArea.tag_remove("bold","sel.first","sel.last")
    else:
        TextArea.tag_add("bold","sel.first","sel.last")

#crete a function for italic text 
def italic_it():
    #creat our font 
    italic_font= font.Font(TextArea,TextArea.cget("font"))
    italic_font.config(slant="italic")

    #configur a tag 
    TextArea.tag_configure("italic",font=italic_font)

    #define current tag
    current_tag = TextArea.tag_names("sel.first")
    #if stetment to see if tag has been bold 
    if "italic" in current_tag:
        TextArea.tag_remove("italic","sel.first","sel.last")
    else:
        TextArea.tag_add("italic","sel.first","sel.last")

#create function for underline
def underline_it():
    #creat our font 
    underline_font= font.Font(TextArea,TextArea.cget("font"))
    underline_font.config(underline=1)

    #configur a tag 
    TextArea.tag_configure("underline",font=underline_font)

    #define current tag
    current_tag = TextArea.tag_names("sel.first")
    #if stetment to see if tag has been bold 
    if "underline" in current_tag:
        TextArea.tag_remove("underline","sel.first","sel.last")
    else:
        TextArea.tag_add("underline","sel.first","sel.last")

#creat a function for selected text color 
def selected_text_color():
    # choose color 
    my_color=colorchooser.askcolor()[1]
    if my_color: 
        #creat our font 
        color_font= font.Font(TextArea,TextArea.cget("font"))
        #configur a tag 
        TextArea.tag_configure("colored",font=color_font,foreground=my_color)
        #define current tag
        current_tag = TextArea.tag_names("sel.first")
        #if stetment to see if tag has been bold 
        if "colored" in current_tag:
            TextArea.tag_remove("colored","sel.first","sel.last")
        else:
            TextArea.tag_add("colored","sel.first","sel.last")

#creat a function for All text color 
def all_text_color():
    my_color=colorchooser.askcolor()[1]
    if my_color:
        TextArea.config(foreground=my_color)

#creat a function for background color
def background_color():
    my_color=colorchooser.askcolor()[1]
    if my_color:
        TextArea.config(background=my_color)

#creat function for font type
def font_chooser(e):
    #choice = clicked_font.get()
    #to_be_change = choice
    TextArea.configure(font=(clicked_font.get(),clicked_size.get()))

#creat function for font size
def font_size_chooser(e):
    # our_font.config(size=drop_size.get(drop_size.curselection()))
    #our_font.config(size=clicked_size.get(drop_size.curselection()))
    TextArea.config(font=(clicked_font.get(),clicked_size.get()))

#creat function for font style 
def font_style_chooser(e):
    style= clicked_style.get().lower()

    if style=="bold":
        TextArea.config(font=(clicked_font.get(),clicked_size.get(),style))#3rd atribute bold in font 
    if style=="italic":
        i=font.Font(family=clicked_font.get(),size=clicked_size.get(),slant="italic")
        TextArea.config(font=i)        
    if style=="regular":
        r=font.Font(family=clicked_font.get(),size=clicked_size.get(),weight="normal",slant="roman",underline=0,overstrike=0)
        TextArea.config(font=r)
    if style=="bold/italic":
        bi=font.Font(family=clicked_font.get(),size=clicked_size.get(),weight="bold",slant="italic")
        TextArea.config(font=bi)
    if style=="underline":
        u=font.Font(family=clicked_font.get(),size=clicked_size.get(),underline=1)
        TextArea.config(font=u)
    if style=="strike":
        o=font.Font(family=clicked_font.get(),size=clicked_size.get(),overstrike=1)
        TextArea.config(font=o)

#creat function for find 
def find_it():
    findentry.focus_set()
    #remove tag found from index 1 to End 
    TextArea.tag_remove('found','1.0',END)

    #get find Entry
    find_get=findentry.get()
    if(find_get):
        idx='1.0'
        while 1:
            #search for desired string from index 1
            idx=TextArea.search(find_get,idx,nocase=1,stopindex=END)

            if not idx: break

            #last index sum of current index and lenth of text
            lastidx='% s+% dc'% (idx,len(find_get)) 

            TextArea.tag_add('found',idx,lastidx)
            idx = lastidx

        #mark located string as red ṇ

        TextArea.tag_config('found',foreground="red")

    findentry.focus_set()

#creat function for replace 
def replace_it():
    replaceentry.focus_set()
    #remove tag found from index 1 to End 
    TextArea.tag_remove('found','1.0',END)

    #get find Entry
    find_get=findentry.get()
    replace_get=replaceentry.get()

    if(find_get and replace_get):
        idx='1.0'

        while 1:
            #search for desired string from index 1
            idx=TextArea.search(find_get,idx,nocase=1,stopindex=END)
            print(idx)
            if not idx :break

            #last index sum of current index and lenth of text
            lastidx='% s+% dc'% (idx,len(find_get)) 

            TextArea.delete(idx,lastidx)
            TextArea.insert(idx,replace_get)

            lastidx='% s+% dc'% (idx,len(replace_get)) 

            #overwrite 'found'at idx
            TextArea.tag_add('found',idx,lastidx)
            idx=lastidx
        # mark laocated string as red 

        TextArea.config('found',foreground="black")
    findentry.focus_set()


#creat function for iserting image 
def insert_image():
    global img
    # Select the imagename from folder
    open_image = filedialog.askopenfilename(title='"pen',filetypes=[("png",".png"),("jpg",".jpg")])
    #open image 
    img = Image.open(open_image)
    #resize the imageṁṇ
    img=img.resize((1200,500),Image.ANTIALIAS)
    img=PhotoImage(file = open_image)

    position= TextArea.index(INSERT)
    TextArea.image_create(position,image=img)

# our font 

our_font=font.Font(family="Helvetica",size="16")

#toolbar frame 
toolbar_frame=Frame(root)
toolbar_frame.pack(padx=5,fill=X)

#creat taxtbox in my_frame 
TextArea=Text(root,font=("helvetica",16),selectbackground="blue",selectforeground="black",undo=True)
TextArea.pack( expand=True,fill=BOTH,padx=5,pady=5)
TextArea.focus_set()


#creat a scrollbar for text area 
text_scroll=Scrollbar(TextArea)
text_scroll.pack(side=RIGHT ,fill=Y)
text_scroll.config(command=TextArea.yview) 
TextArea.config(yscrollcommand=text_scroll.set)

#creat statusbar
status_bar=Label(root,text="ready        ",anchor=E)
status_bar.pack(fill=X,side=BOTTOM,ipadx=5)

#create menu bar 
mainmenu=Menu(root)
#create sub menu in main menu 
file_menu=Menu(mainmenu,tearoff=0)
file_menu.add_command(label="New",command=new_f)
#file_menu.add_command(label="Open PDF",command=open_pdf)
#file_menu.add_command(label="Open DOCX",command=open_docx)
file_menu.add_command(label="Open",command=open_f)
file_menu.add_command(label="Save",command=save_f)
file_menu.add_command(label="Save As",command=save_as_f)
file_menu.add_command(label="Save As Docx",command=save_as_docx)
file_menu.add_separator()
file_menu.add_command(label="exit",command=root.quit)
mainmenu.add_cascade(label="File",menu=file_menu)

edit_menu=Menu(mainmenu,tearoff=0)
edit_menu.add_command(label="Undo",command=TextArea.edit_undo,accelerator="(Ctrl+z)")
edit_menu.add_command(label="Redo",command=TextArea.edit_redo,accelerator="(Ctrl+y)")
edit_menu.add_separator()
edit_menu.add_command(label="Cut",command=cut_f,accelerator="(Ctrl+x)")
edit_menu.add_command(label="Copy",command=copy_f,accelerator="(Ctrl+c)")
edit_menu.add_command(label="Paste",command=paste_f,accelerator="(Ctrl+v)")
edit_menu.add_separator()
edit_menu.add_command(label="Select All",command=lambda: select_all_f(True),accelerator="(Ctrl+a)")
edit_menu.add_command(label="Clear",command=clear_f)
edit_menu.add_separator()
edit_menu.add_command(label="Find",command=find_it)
edit_menu.add_command(label="Replace",command=replace_it)
mainmenu.add_cascade(label="Edit",menu=edit_menu)

formate_menu=Menu(mainmenu,tearoff=0)
formate_menu.add_command(label="Bold",command=bold_it)  
formate_menu.add_command(label="Italic",command=italic_it)  
formate_menu.add_command(label="Underline",command=underline_it)  
formate_menu.add_separator() 
#create sub menu of formate menu 
color_menu=Menu(formate_menu,tearoff=0)
color_menu.add_command(label="Selected text",command=selected_text_color)
color_menu.add_command(label="All text",command=all_text_color)
formate_menu.add_cascade(label="Font Color",menu=color_menu)

formate_menu.add_command(label="Background Color",command=background_color)  
mainmenu.add_cascade(label="Formate",menu=formate_menu)

help_menu=Menu(mainmenu,tearoff=0)
help_menu.add_command(label="About",command=about_f)
mainmenu.add_cascade(label="Help",menu=help_menu)

root.config(menu=mainmenu)

#button for toolbar

#bold button
bold_button=Button(toolbar_frame,text="Bold",command=bold_it)
bold_button.grid(row=0,column=0,pady=5,padx=2,sticky=W)
#italic button
italic_button=Button(toolbar_frame,text="Italic",command=italic_it)
italic_button.grid(row=0,column=1,padx=2,pady=5,sticky=W)
# under line button
underline_button=Button(toolbar_frame,text="UnderLine",command=underline_it)
underline_button.grid(row=0,column=2,padx=2,pady=5,sticky=W)


#create lable & list  for font related things in toolbar frame 



# lable for font style 
font_label=Label(toolbar_frame,text="Font Type:",font=("Helvetica",14))
font_label.grid(row=0,column=3,padx=10,sticky=W)
#lable for font size 
font_size_label=Label(toolbar_frame,text="Font Size:",font=("Helvetica",14))
font_size_label.grid(row=0,column=4,padx=10,sticky=W)
#lable for font style
font_style_label=Label(toolbar_frame,text="Font Style:",font=("Helvetica",14))
font_style_label.grid(row=0,column=5,padx=10,sticky=W) 

#create lists for font

# drop down menu for font
# font_list=[]
# for f in font.families():
#     font_list.insert('end',f)
font_list = ['Helvetica','Garamond','CircusDog','Dannette','DraftHand','Flowerport','Grimmy','HolyCow','Isepic','JohnDoe','Keener','Network','Orion','Pesto','ReadOut','Rockwell','SmithPremier','Torcho','WST_Engl','WST_Fren','ZipSonik'] 
clicked_font=StringVar()
clicked_font.set("helvetica")
drop_font=OptionMenu(toolbar_frame, clicked_font, *font_list,command = font_chooser)
drop_font.grid(row=1,column=3,padx=5,pady=5)

# drop down menu for font size
size_list=[8,10,12,14,16,18,20,24,28,32,34,36,40,48,60,72]
clicked_size=IntVar()
clicked_size.set(16)
drop_size=OptionMenu(toolbar_frame,clicked_size,*size_list,command= font_size_chooser)
drop_size.grid(row=1,column=4,padx=5,pady=5)

# drop down menu for font style
style_list=["regular","Bold","Italic","Bold/Italic","Underline","Strike"]
clicked_style=StringVar()
clicked_style.set("regular");
drop_style=OptionMenu(toolbar_frame,clicked_style,*style_list,command=font_style_chooser)
drop_style.grid(row=1,column=5,padx=5,pady=5)

# find lable entry and button button 
Label(toolbar_frame,text="Find:",font=("Helvetica",8)).grid(row=0,column=6,padx=5,pady=5)
#findvalue=str()
findentry=Entry(toolbar_frame)
findentry.grid(row=1,column=6,padx=10,pady=5)
#findentry.focus_set()
Button(toolbar_frame,text="Find",command=find_it).grid(row=3,column=6,padx=10,pady=5)

# replace lable entry and button button 
Label(toolbar_frame,text="Replace with:",font=("Helvetica",8)).grid(row=0,column=7,padx=10,pady=5)
#replacevalue=str()
replaceentry=Entry(toolbar_frame)
replaceentry.grid(row=1,column=7,padx=10,pady=5)
#replaceentry.focus_set()
Button(toolbar_frame,text="Replace",command=replace_it).grid(row=3,column=7,padx=10,pady=5)

#creat image insert button 
insert_image_button=Button(toolbar_frame,text="Insert Image",command=insert_image)
insert_image_button.grid(row=0,column=8,padx=10,pady=5)             

#binding 
#bind ctrl+A to select all of edit sub menu 
root.bind('<Control-A>',select_all_f)
root.bind('<Control-a>',select_all_f)

root.mainloop()